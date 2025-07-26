# TranscriptionThread クラスを拡張して実際のWhisperモデルを使用するバージョン
import whisper
import torch
import os
import tempfile
from PyQt5.QtCore import QThread, pyqtSignal

class WhisperTranscriptionThread(QThread):
    """Whisperモデルを使用した音声文字起こし処理を行うスレッド"""
    progress_signal = pyqtSignal(int)
    log_signal = pyqtSignal(str)
    finished_signal = pyqtSignal(str, str)  # ファイル名、テキスト内容

    def __init__(self, file_path, language='ja', model_size='base'):
        super().__init__()
        self.file_path = file_path
        self.language = language
        self.model_size = model_size
        self.model = None
        
    def run(self):
        file_name = os.path.basename(self.file_path)
        self.log_signal.emit(f"処理開始: {file_name}")
        
        try:
            # モデルが初期化されていない場合は初期化
            if self.model is None:
                self.log_signal.emit(f"Whisperモデル '{self.model_size}' をロード中...")
                self.progress_signal.emit(10)
                
                # GPUが利用可能であれば使用
                device = "cuda" if torch.cuda.is_available() else "cpu"
                self.log_signal.emit(f"使用デバイス: {device}")
                
                # モデルのロード
                self.model = whisper.load_model(self.model_size, device=device)
                self.log_signal.emit("モデルロード完了")
                self.progress_signal.emit(30)
            
            # 言語設定
            language_code = self.language if self.language != 'auto' else None
            
            # 音声認識のオプション
            options = {
                "language": language_code,  # 言語を指定（Noneの場合は自動検出）
                "task": "transcribe",       # 文字起こしタスク
                "verbose": False            # 詳細ログは無効
            }
            
            self.log_signal.emit(f"音声認識処理中: {file_name}...")
            self.progress_signal.emit(40)
            
            # 音声認識実行
            result = self.model.transcribe(self.file_path, **options)
            self.progress_signal.emit(90)
            
            # 結果の取得
            transcribed_text = result["text"]
            detected_language = result.get("language", "不明")
            
            # メタデータを含めた出力テキストの作成
            output_text = f"# 文字起こし結果: {file_name}\n\n"
            output_text += f"言語: {detected_language}\n"
            output_text += f"モデル: {self.model_size}\n\n"
            output_text += "## テキスト内容\n\n"
            output_text += transcribed_text
            
            self.log_signal.emit(f"処理完了: {file_name} ({detected_language})")
            self.progress_signal.emit(100)
            self.finished_signal.emit(file_name, output_text)
            
        except Exception as e:
            self.log_signal.emit(f"エラー: {file_name} - {str(e)}")


# メインアプリケーションへの統合例 (start_transcription メソッドの一部を更新)
def start_transcription(self):
    """文字起こし処理を開始"""
    if not self.selected_files:
        self.log_text.append("ファイルが選択されていません。")
        return
    
    # 既存のスレッドをクリア
    for thread in self.active_threads:
        if thread.isRunning():
            thread.terminate()
    
    self.active_threads = []
    self.transcription_results = {}
    self.progress_bar.setValue(0)
    
    # 言語設定の取得
    language_map = {
        "日本語": "ja", 
        "英語": "en", 
        "中国語": "zh", 
        "韓国語": "ko", 
        "自動検出": "auto"
    }
    selected_language = language_map[self.language_combo.currentText()]
    
    # モデルサイズの取得 (UIに追加が必要)
    model_size = "base"  # tiny, base, small, medium, large から選択
    
    # 最初のファイルの処理を開始
    self.start_next_file(0, selected_language, model_size)

def start_next_file(self, index, language, model_size):
    """次のファイルの処理を開始"""
    if index < len(self.selected_files):
        file_path = self.selected_files[index]
        
        # WhisperTranscriptionThread を使用
        thread = WhisperTranscriptionThread(file_path, language, model_size)
        thread.progress_signal.connect(self.update_progress)
        thread.log_signal.connect(self.update_log)
        thread.finished_signal.connect(
            lambda file_name, text: self.handle_transcription_finished(
                file_name, text, index, language, model_size
            )
        )
        
        self.active_threads.append(thread)
        thread.start()

def handle_transcription_finished(self, file_name, text, current_index, language, model_size):
    """文字起こし完了時の処理"""
    # 結果を保存
    self.transcription_results[file_name] = text
    
    # 出力形式に基づいたファイル保存
    format_map = {
        "テキストファイル (.txt)": ".txt",
        "Word文書 (.docx)": ".docx",
        "JSONファイル (.json)": ".json"
    }
    
    selected_format = format_map[self.format_combo.currentText()]
    base_name = os.path.splitext(file_name)[0]
    
    if self.output_dir:
        output_path = os.path.join(self.output_dir, f"{base_name}{selected_format}")
    else:
        output_path = f"{base_name}{selected_format}"
    
    try:
        if selected_format == ".txt":
            # テキストファイルとして保存
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
        elif selected_format == ".docx":
            # Word文書として保存 (python-docxライブラリが必要)
            try:
                from docx import Document
                document = Document()
                document.add_heading(f'文字起こし結果: {file_name}', 0)
                
                # テキストの処理 (メタデータとコンテンツを分離)
                lines = text.split('\n')
                metadata_end = 0
                for i, line in enumerate(lines):
                    if line.startswith('## テキスト内容'):
                        metadata_end = i
                        break
                
                # メタデータ
                for i in range(1, metadata_end):
                    if lines[i].strip():
                        key, value = lines[i].split(':', 1)
                        document.add_paragraph(f"{key}: {value.strip()}")
                
                # 本文テキスト
                document.add_heading('テキスト内容', level=1)
                document.add_paragraph(
                    '\n'.join(lines[metadata_end+1:])
                )
                
                document.save(output_path)
            except ImportError:
                self.log_text.append("エラー: python-docxライブラリがインストールされていません。テキスト形式で保存します。")
                with open(output_path.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
                    f.write(text)
        elif selected_format == ".json":
            # JSON形式で保存
            import json
            
            # メタデータとテキスト内容を分離
            metadata = {}
            content = ""
            in_content = False
            
            for line in text.split('\n'):
                if line.startswith('## テキスト内容'):
                    in_content = True
                    continue
                
                if not in_content and ':' in line:
                    key, value = line.split(':', 1)
                    metadata[key.strip()] = value.strip()
                elif in_content:
                    content += line + '\n'
            
            json_data = {
                "filename": file_name,
                "metadata": metadata,
                "content": content.strip()
            }
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, ensure_ascii=False, indent=2)
                
        self.log_text.append(f"保存完了: {output_path}")
    except Exception as e:
        self.log_text.append(f"ファイル保存エラー: {str(e)}")
    
    # 次のファイルを処理
    next_index = current_index + 1
    if next_index < len(self.selected_files):
        self.start_next_file(next_index, language, model_size)
    else:
        self.log_text.append("全ファイルの処理が完了しました。")
        # 処理完了通知
        self.progress_bar.setValue(100)
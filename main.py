import sys
import os
import logging
import traceback
from datetime import datetime
from PyQt5.QtWidgets import (QApplication, QMainWindow, QPushButton, QVBoxLayout, QHBoxLayout, 
                             QWidget, QFileDialog, QListWidget, QProgressBar, QLabel, 
                             QTextEdit, QComboBox, QGroupBox, QGridLayout, QCheckBox, QMessageBox)
from PyQt5.QtCore import Qt, QThread, pyqtSignal

# ロギングの設定
log_directory = "logs"
if not os.path.exists(log_directory):
    os.makedirs(log_directory)

log_filename = os.path.join(log_directory, f"transcriber_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")

# ロガーの設定
logger = logging.getLogger("MP3Transcriber")
logger.setLevel(logging.DEBUG)

# ファイルハンドラ
file_handler = logging.FileHandler(log_filename, encoding='utf-8')
file_handler.setLevel(logging.DEBUG)

# コンソールハンドラ
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)

# フォーマッタ
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
file_handler.setFormatter(formatter)
console_handler.setFormatter(formatter)

# ハンドラの追加
logger.addHandler(file_handler)
logger.addHandler(console_handler)

# グローバルな例外ハンドラ
def exception_hook(exc_type, exc_value, exc_traceback):
    logger.critical("Uncaught exception", exc_info=(exc_type, exc_value, exc_traceback))
    sys.__excepthook__(exc_type, exc_value, exc_traceback)

sys.excepthook = exception_hook

# Whisperモデルを使用した音声文字起こしスレッド
class WhisperTranscriptionThread(QThread):
    """Whisperモデルを使用した音声文字起こし処理を行うスレッド"""
    progress_signal = pyqtSignal(int)
    log_signal = pyqtSignal(str)
    finished_signal = pyqtSignal(str, str)  # ファイル名、テキスト内容
    error_signal = pyqtSignal(str, str)  # エラーメッセージ、詳細

    def __init__(self, file_path, language='ja', model_size='base'):
        super().__init__()
        self.file_path = file_path
        self.language = language
        self.model_size = model_size
        self.model = None
        
    def run(self):
        file_name = os.path.basename(self.file_path)
        logger.info(f"処理開始: {file_name}")
        self.log_signal.emit(f"処理開始: {file_name}")
        
        try:
            # モデルが初期化されていない場合は初期化
            if self.model is None:
                logger.info(f"Whisperモデル '{self.model_size}' をロード中...")
                self.log_signal.emit(f"Whisperモデル '{self.model_size}' をロード中...")
                self.progress_signal.emit(10)
                # Whisperモジュールのインポート
                try:
                    logger.debug("Whisperモジュールをインポート中...")
                    import whisper
                    import torch
                    logger.debug("Whisperモジュールのインポート成功")
                except ImportError as e:
                    error_msg = f"必要なライブラリがインストールされていません: {str(e)}"
                    logger.error(error_msg)
                    self.log_signal.emit(error_msg)
                    self.log_signal.emit("以下のコマンドを実行してください: uv pip install openai-whisper torch")
                    self.error_signal.emit("ライブラリエラー", error_msg)
                    return
                
                # GPUが利用可能であれば使用
                device = "cuda" if torch.cuda.is_available() else "cpu"
                logger.info(f"使用デバイス: {device}")
                self.log_signal.emit(f"使用デバイス: {device}")
                
                try:
                    # モデルのロード
                    logger.debug(f"モデル {self.model_size} をロード中...")
                    self.model = whisper.load_model(self.model_size, device=device)
                    logger.info("モデルロード完了")
                    self.log_signal.emit("モデルロード完了")
                    self.progress_signal.emit(30)
                except Exception as e:
                    error_msg = f"モデルのロードに失敗しました: {str(e)}"
                    logger.error(error_msg)
                    logger.error(traceback.format_exc())
                    self.log_signal.emit(error_msg)
                    self.error_signal.emit("モデルロードエラー", traceback.format_exc())
                    return
            
            # 言語設定
            language_code = self.language if self.language != 'auto' else None
            logger.debug(f"言語設定: {language_code}")
            
            # 音声認識のオプション
            options = {
                "language": language_code,  # 言語を指定（Noneの場合は自動検出）
                "task": "transcribe",       # 文字起こしタスク
                "verbose": False            # 詳細ログは無効
            }
            
            logger.info(f"音声認識処理中: {file_name}...")
            self.log_signal.emit(f"音声認識処理中: {file_name}...")
            self.progress_signal.emit(40)
            
            try:
                # 音声ファイルの存在確認
                if not os.path.exists(self.file_path):
                    error_msg = f"ファイルが見つかりません: {self.file_path}"
                    logger.error(error_msg)
                    self.log_signal.emit(error_msg)
                    self.error_signal.emit("ファイルエラー", error_msg)
                    return
                
                # ファイルサイズのログ
                file_size = os.path.getsize(self.file_path) / (1024 * 1024)  # MB単位
                logger.debug(f"ファイルサイズ: {file_size:.2f} MB")
                
                # 音声認識実行
                logger.debug(f"Whisperで音声認識を実行中: {self.file_path}")
                result = self.model.transcribe(self.file_path, **options)
                logger.debug("音声認識完了")
                self.progress_signal.emit(90)
                
                # 結果の取得
                logger.debug("音声認識結果を取得中")
                transcribed_text = result["text"]
                detected_language = result.get("language", "不明")
                logger.debug(f"検出された言語: {detected_language}")
                logger.debug(f"テキスト長: {len(transcribed_text)} 文字")
                
                # メタデータを含めた出力テキストの作成
                output_text = f"# 文字起こし結果: {file_name}\n\n"
                output_text += f"言語: {detected_language}\n"
                output_text += f"モデル: {self.model_size}\n\n"
                output_text += "## テキスト内容\n\n"
                output_text += transcribed_text
                
                logger.info(f"処理完了: {file_name} ({detected_language})")
                self.log_signal.emit(f"処理完了: {file_name} ({detected_language})")
                self.progress_signal.emit(100)
                self.finished_signal.emit(file_name, output_text)
                
            except Exception as e:
                error_msg = f"音声認識処理でエラーが発生しました: {str(e)}"
                logger.error(error_msg)
                logger.error(traceback.format_exc())
                self.log_signal.emit(error_msg)
                self.error_signal.emit("音声認識エラー", traceback.format_exc())
            
        except Exception as e:
            error_msg = f"エラー: {file_name} - {str(e)}"
            logger.error(error_msg)
            logger.error(traceback.format_exc())
            self.log_signal.emit(error_msg)
            self.error_signal.emit("一般エラー", traceback.format_exc())


class MP3TranscriberApp(QMainWindow):
    """MP3文字起こしアプリケーションのメインウィンドウ"""
    
    def __init__(self):
        super().__init__()
        self.setWindowTitle("MP3文字起こしアプリ")
        self.setGeometry(100, 100, 800, 600)
        self.selected_files = []
        self.output_dir = ""
        self.active_threads = []
        self.transcription_results = {}  # ファイル名:テキスト内容
        self.whisper_model = None  # Whisperモデルのインスタンス
        
        logger.info("アプリケーション初期化開始")
        self.init_ui()
        logger.info("アプリケーション初期化完了")
        
    def init_ui(self):
        logger.debug("UI初期化開始")
        # メインレイアウト
        main_layout = QVBoxLayout()
        
        # ファイル選択エリア
        file_group = QGroupBox("ファイル選択")
        file_layout = QVBoxLayout()
        
        browse_layout = QHBoxLayout()
        self.folder_btn = QPushButton("フォルダ選択")
        self.folder_btn.clicked.connect(self.select_folder)
        self.files_btn = QPushButton("ファイル選択")
        self.files_btn.clicked.connect(self.select_files)
        browse_layout.addWidget(self.folder_btn)
        browse_layout.addWidget(self.files_btn)
        
        self.file_list = QListWidget()
        
        file_layout.addLayout(browse_layout)
        file_layout.addWidget(QLabel("選択されたファイル:"))
        file_layout.addWidget(self.file_list)
        
        file_group.setLayout(file_layout)
        
        # 設定エリア
        settings_group = QGroupBox("設定")
        settings_layout = QGridLayout()
        
        settings_layout.addWidget(QLabel("言語:"), 0, 0)
        self.language_combo = QComboBox()
        self.language_combo.addItems(["日本語", "英語", "中国語", "韓国語", "自動検出"])
        settings_layout.addWidget(self.language_combo, 0, 1)
        
        settings_layout.addWidget(QLabel("モデルサイズ:"), 0, 2)
        self.model_combo = QComboBox()
        self.model_combo.addItems(["tiny", "base", "small", "medium", "large"])
        self.model_combo.setCurrentText("base")  # デフォルトはbaseモデル
        settings_layout.addWidget(self.model_combo, 0, 3)
        
        settings_layout.addWidget(QLabel("出力先:"), 1, 0)
        output_layout = QHBoxLayout()
        self.output_path_label = QLabel("デフォルト: カレントディレクトリ")
        self.output_browse_btn = QPushButton("参照...")
        self.output_browse_btn.clicked.connect(self.select_output_dir)
        output_layout.addWidget(self.output_path_label)
        output_layout.addWidget(self.output_browse_btn)
        settings_layout.addLayout(output_layout, 1, 1, 1, 3)
        
        settings_layout.addWidget(QLabel("出力形式:"), 2, 0)
        self.format_combo = QComboBox()
        self.format_combo.addItems(["テキストファイル (.txt)", "Word文書 (.docx)", "JSONファイル (.json)"])
        settings_layout.addWidget(self.format_combo, 2, 1, 1, 3)
        
        # デバッグモード
        debug_layout = QHBoxLayout()
        self.debug_checkbox = QCheckBox("デバッグモード")
        self.debug_checkbox.setChecked(True)  # デフォルトでオン
        debug_layout.addWidget(self.debug_checkbox)
        settings_layout.addLayout(debug_layout, 3, 0, 1, 4)
        
        settings_group.setLayout(settings_layout)
        
        # 処理エリア
        process_group = QGroupBox("処理")
        process_layout = QVBoxLayout()
        
        button_layout = QHBoxLayout()
        self.start_btn = QPushButton("文字起こし開始")
        self.start_btn.clicked.connect(self.start_transcription)
        self.start_btn.setEnabled(False)
        
        self.cancel_btn = QPushButton("処理中止")
        self.cancel_btn.clicked.connect(self.cancel_transcription)
        self.cancel_btn.setEnabled(False)
        
        button_layout.addWidget(self.start_btn)
        button_layout.addWidget(self.cancel_btn)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        
        process_layout.addLayout(button_layout)
        process_layout.addWidget(QLabel("進捗状況:"))
        process_layout.addWidget(self.progress_bar)
        process_layout.addWidget(QLabel("ログ:"))
        process_layout.addWidget(self.log_text)
        
        process_group.setLayout(process_layout)
        
        # レイアウトの組み立て
        main_layout.addWidget(file_group)
        main_layout.addWidget(settings_group)
        main_layout.addWidget(process_group)
        
        # 中央ウィジェットの設定
        central_widget = QWidget()
        central_widget.setLayout(main_layout)
        self.setCentralWidget(central_widget)
        
        # 初期ログメッセージ
        self.log_text.append(f"MP3文字起こしアプリを起動しました。ログファイル: {log_filename}")
        self.log_text.append("フォルダまたはファイルを選択してください。")
        logger.debug("UI初期化完了")

    def select_folder(self):
        """フォルダを選択し、MP3ファイルを検索"""
        logger.debug("フォルダ選択ダイアログを開始")
        folder_path = QFileDialog.getExistingDirectory(self, "フォルダ選択", "")
        if folder_path:
            logger.info(f"選択されたフォルダ: {folder_path}")
            self.file_list.clear()
            self.selected_files = []
            
            try:
                for root, _, files in os.walk(folder_path):
                    for file in files:
                        if file.lower().endswith('.mp3'):
                            file_path = os.path.join(root, file)
                            self.selected_files.append(file_path)
                            self.file_list.addItem(file)
                
                if self.selected_files:
                    self.start_btn.setEnabled(True)
                    logger.info(f"{len(self.selected_files)}個のMP3ファイルが見つかりました")
                    self.log_text.append(f"{len(self.selected_files)}個のMP3ファイルが見つかりました。")
                else:
                    self.start_btn.setEnabled(False)
                    logger.warning(f"選択されたフォルダ内にMP3ファイルが見つかりませんでした: {folder_path}")
                    self.log_text.append("MP3ファイルが見つかりませんでした。")
            except Exception as e:
                logger.error(f"フォルダ走査中にエラーが発生しました: {str(e)}")
                logger.error(traceback.format_exc())
                self.log_text.append(f"エラー: フォルダの走査に失敗しました - {str(e)}")
    
    def select_files(self):
        """複数のMP3ファイルを選択"""
        logger.debug("ファイル選択ダイアログを開始")
        files, _ = QFileDialog.getOpenFileNames(self, "MP3ファイル選択", "", "MP3 Files (*.mp3)")
        if files:
            logger.info(f"{len(files)}個のファイルが選択されました")
            self.file_list.clear()
            self.selected_files = files
            
            for file in files:
                logger.debug(f"選択されたファイル: {file}")
                self.file_list.addItem(os.path.basename(file))
            
            self.start_btn.setEnabled(True)
            self.log_text.append(f"{len(files)}個のファイルが選択されました。")
    
    def select_output_dir(self):
        """出力ディレクトリを選択"""
        logger.debug("出力先選択ダイアログを開始")
        dir_path = QFileDialog.getExistingDirectory(self, "出力先フォルダ選択", "")
        if dir_path:
            logger.info(f"出力先フォルダ: {dir_path}")
            self.output_dir = dir_path
            self.output_path_label.setText(dir_path)
            self.log_text.append(f"出力先: {dir_path}")
    
    def start_transcription(self):
        """文字起こし処理を開始"""
        if not self.selected_files:
            logger.warning("ファイルが選択されていません")
            self.log_text.append("ファイルが選択されていません。")
            return
        
        # 既存のスレッドをクリア
        for thread in self.active_threads:
            if thread.isRunning():
                thread.terminate()
        
        self.active_threads = []
        self.transcription_results = {}
        self.progress_bar.setValue(0)
        
        # UI状態の更新
        self.start_btn.setEnabled(False)
        self.cancel_btn.setEnabled(True)
        self.folder_btn.setEnabled(False)
        self.files_btn.setEnabled(False)
        
        # 言語設定の取得
        language_map = {
            "日本語": "ja", 
            "英語": "en", 
            "中国語": "zh", 
            "韓国語": "ko", 
            "自動検出": "auto"
        }
        selected_language = language_map[self.language_combo.currentText()]
        
        # モデルサイズの取得
        model_size = self.model_combo.currentText()
        
        # デバッグモード確認
        debug_mode = self.debug_checkbox.isChecked()
        if debug_mode:
            logger.setLevel(logging.DEBUG)
            logger.debug("デバッグモード有効")
        else:
            logger.setLevel(logging.INFO)
        
        # ログに設定情報を表示
        logger.info(f"言語設定: {self.language_combo.currentText()} ({selected_language})")
        logger.info(f"使用モデル: {model_size}")
        logger.info(f"出力形式: {self.format_combo.currentText()}")
        
        self.log_text.append(f"言語設定: {self.language_combo.currentText()}")
        self.log_text.append(f"使用モデル: {model_size}")
        self.log_text.append(f"出力形式: {self.format_combo.currentText()}")
        self.log_text.append(f"デバッグモード: {'有効' if debug_mode else '無効'}")
        
        # 最初のファイルの処理を開始
        self.start_next_file(0, selected_language, model_size)
    
    def cancel_transcription(self):
        """処理中の文字起こしをキャンセル"""
        logger.info("処理中止リクエスト")
        for thread in self.active_threads:
            if thread.isRunning():
                thread.terminate()
        
        logger.info("処理を中止しました")
        self.log_text.append("処理を中止しました。")
        
        # UI状態の更新
        self.start_btn.setEnabled(True)
        self.cancel_btn.setEnabled(False)
        self.folder_btn.setEnabled(True)
        self.files_btn.setEnabled(True)
    
    def start_next_file(self, index, language, model_size):
        """次のファイルの処理を開始"""
        if index < len(self.selected_files):
            file_path = self.selected_files[index]
            file_name = os.path.basename(file_path)
            
            logger.info(f"{index+1}/{len(self.selected_files)}: {file_name} の処理を開始します")
            self.log_text.append(f"{index+1}/{len(self.selected_files)}: {file_name} の処理を開始します...")
            
            # WhisperTranscriptionThread を使用
            thread = WhisperTranscriptionThread(file_path, language, model_size)
            thread.progress_signal.connect(self.update_progress)
            thread.log_signal.connect(self.update_log)
            thread.error_signal.connect(self.handle_error)
            thread.finished_signal.connect(
                lambda file_name, text: self.handle_transcription_finished(
                    file_name, text, index, language, model_size
                )
            )
            
            self.active_threads.append(thread)
            thread.start()
            logger.debug(f"スレッド開始: {file_name}")
        else:
            # 全ファイルの処理完了
            logger.info("全ファイルの処理が完了しました")
            self.log_text.append("全ファイルの処理が完了しました。")
            
            # UI状態の更新
            self.start_btn.setEnabled(True)
            self.cancel_btn.setEnabled(False)
            self.folder_btn.setEnabled(True)
            self.files_btn.setEnabled(True)
    
    def update_progress(self, value):
        """進捗バーを更新"""
        self.progress_bar.setValue(value)
    
    def update_log(self, message):
        """ログを更新"""
        self.log_text.append(message)
        # 自動スクロール
        self.log_text.verticalScrollBar().setValue(self.log_text.verticalScrollBar().maximum())
    
    def handle_error(self, title, message):
        """エラーメッセージを表示"""
        logger.error(f"エラー: {title} - {message}")
        QMessageBox.critical(self, f"エラー: {title}", message)
    
    def handle_transcription_finished(self, file_name, text, current_index, language, model_size):
        """文字起こし完了時の処理"""
        logger.debug(f"文字起こし完了: {file_name}")
        
        # 結果を保存
        self.transcription_results[file_name] = text
        
        # 出力形式に基づいたファイル保存
        format_map = {
            "テキストファイル (.txt)": ".txt",
            "Word文書 (.docx)": ".docx",
            "JSONファイル (.json)": ".json"
        }
        
        selected_format = format_map[self.format_combo.currentText()]
        base_name = os.path.splitext(file_name)[0]
        
        if self.output_dir:
            output_path = os.path.join(self.output_dir, f"{base_name}{selected_format}")
        else:
            output_path = f"{base_name}{selected_format}"
        
        logger.debug(f"保存先: {output_path}")
        
        try:
            if selected_format == ".txt":
                # テキストファイルとして保存
                logger.debug(f"テキストファイルを保存中: {output_path}")
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                logger.debug("テキストファイル保存完了")
            elif selected_format == ".docx":
                # Word文書として保存 (python-docxライブラリが必要)
                try:
                    logger.debug("Word文書として保存中")
                    from docx import Document
                    document = Document()
                    document.add_heading(f'文字起こし結果: {file_name}', 0)
                    
                    # テキストの処理 (メタデータとコンテンツを分離)
                    lines = text.split('\n')
                    metadata_end = 0
                    for i, line in enumerate(lines):
                        if line.startswith('## テキスト内容'):
                            metadata_end = i
                            break
                    
                    # メタデータ
                    for i in range(1, metadata_end):
                        if lines[i].strip():
                            key, value = lines[i].split(':', 1)
                            document.add_paragraph(f"{key}: {value.strip()}")
                    
                    # 本文テキスト
                    document.add_heading('テキスト内容', level=1)
                    document.add_paragraph(
                        '\n'.join(lines[metadata_end+1:])
                    )
                    
                    document.save(output_path)
                    logger.debug("Word文書保存完了")
                except ImportError as e:
                    logger.warning(f"python-docxライブラリがインストールされていません: {str(e)}")
                    self.log_text.append("エラー: python-docxライブラリがインストールされていません。テキスト形式で保存します。")
                    with open(output_path.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
                        f.write(text)
            elif selected_format == ".json":
                # JSON形式で保存
                logger.debug("JSONファイルとして保存中")
                import json
                
                # メタデータとテキスト内容を分離
                metadata = {}
                content = ""
                in_content = False
                
                for line in text.split('\n'):
                    if line.startswith('## テキスト内容'):
                        in_content = True
                        continue
                    
                    if not in_content and ':' in line:
                        key, value = line.split(':', 1)
                        metadata[key.strip()] = value.strip()
                    elif in_content:
                        content += line + '\n'
                
                json_data = {
                    "filename": file_name,
                    "metadata": metadata,
                    "content": content.strip()
                }
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(json_data, f, ensure_ascii=False, indent=2)
                logger.debug("JSONファイル保存完了")
                    
            logger.info(f"保存完了: {output_path}")
            self.log_text.append(f"保存完了: {output_path}")
        except Exception as e:
            error_msg = f"ファイル保存エラー: {str(e)}"
            logger.error(error_msg)
            logger.error(traceback.format_exc())
            self.log_text.append(error_msg)
            
            # エラーメッセージを表示
            QMessageBox.warning(self, "保存エラー", f"ファイル保存中にエラーが発生しました:\n{str(e)}")
        
        # 次のファイルを処理
        next_index = current_index + 1
        if next_index < len(self.selected_files):
            self.start_next_file(next_index, language, model_size)
        else:
            # 処理完了通知
            self.progress_bar.setValue(100)
            logger.info("全ファイルの処理が完了しました")
            self.log_text.append("全ファイルの処理が完了しました。")
            
            # UI状態の更新
            self.start_btn.setEnabled(True)
            self.cancel_btn.setEnabled(False)
            self.folder_btn.setEnabled(True)
            self.files_btn.setEnabled(True)


def main():
    try:
        logger.info("アプリケーション起動")
        app = QApplication(sys.argv)
        window = MP3TranscriberApp()
        window.show()
        logger.info("アプリケーションウィンドウを表示")
        sys.exit(app.exec_())
    except Exception as e:
        logger.critical(f"アプリケーション実行中に致命的なエラーが発生しました: {str(e)}")
        logger.critical(traceback.format_exc())


if __name__ == "__main__":
    main()